#!/usr/bin/env python3
"""
Конвертація 3 джерел Академії у структурований JSON для бібліотеки гри.

Запуск:  python3 scripts/convert_library.py

Виходи:  src/data/library/cosmo.json, methodichka.json, united.json,
         + index.json (загальний пошуковий індекс)
"""

import json
import re
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parent.parent
LIBRARY_DIR = PROJECT_ROOT / "src" / "data" / "library"
LIBRARY_DIR.mkdir(parents=True, exist_ok=True)

DESKTOP = Path("/Users/dr.kay.avi/Desktop")
PDF_COSMO = DESKTOP / "ФІнальні книги" / "космоенергетика повна.pdf"
DOCX_METH = DESKTOP / "ФІнальні книги" / "Методична_книга_АКАДЕМІЯ_повна.docx"
DOCX_UNITED = DESKTOP / "модулі навчання" / "Обєднані файли" / "Повна_об'єднана_база_текстів_Космоенергетика_Контроль_Потоки_Код.docx"


def clean_text(s: str) -> str:
    s = s.replace(" ", " ")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()


def chunk_by_headers(paragraphs, header_predicate, source_id):
    """Розбити параграфи на секції за функцією-предикатом заголовка."""
    sections = []
    current = {"id": f"{source_id}-intro", "title": "Передмова", "body": []}
    counter = 0
    for p in paragraphs:
        text = p.strip()
        if not text:
            continue
        if header_predicate(text):
            if current["body"]:
                current["body"] = clean_text("\n".join(current["body"]))
                sections.append(current)
            counter += 1
            slug = re.sub(r"[^a-zа-яёіїє0-9]+", "-", text.lower(), flags=re.IGNORECASE)[:40].strip("-")
            current = {
                "id": f"{source_id}-{counter:02d}-{slug}",
                "title": text,
                "body": [],
            }
        else:
            current["body"].append(text)
    if current["body"]:
        current["body"] = clean_text("\n".join(current["body"]))
        sections.append(current)
    return sections


def is_methodichka_header(s: str) -> bool:
    if len(s) > 90 or len(s) < 4:
        return False
    if s.startswith("•") or s.startswith("⸻"):
        return False
    if s.endswith(":") and len(s) < 60:
        return True
    if re.match(r"^(Розділ|Глава|Том|Частина|Передмова|Епілог|Додаток|Глосарій)\b", s, re.IGNORECASE):
        return True
    if re.match(r"^[А-ЯҐЇЄІA-Z][А-ЯҐЇЄІA-Z .\-]+$", s) and len(s) < 70:
        return True
    return False


def is_united_header(s: str) -> bool:
    if not s or len(s) > 100:
        return False
    return bool(re.match(r"^[\U0001F300-\U0001FAFF✦◯◇♥◐⚡☉☼☾⊹∞◌🜃🜂🜁🜄]\s", s))


def is_cosmo_pdf_header(s: str) -> bool:
    if not s or len(s) > 100 or len(s) < 4:
        return False
    if re.match(r"^(Канал|Блок|Розділ|Частина|Глава)\b", s, re.IGNORECASE):
        return True
    if re.match(r"^(\d{1,2})\.\s+\S", s):
        return True
    if re.match(r"^[А-ЯҐЇЄІ][А-ЯҐЇЄІ\- .]{3,}$", s):
        return True
    return False


def convert_pdf(path: Path, source_id: str, source_label: str) -> dict:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    full_text = []
    for page in reader.pages:
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        if t.strip():
            full_text.append(t)
    raw = "\n".join(full_text)
    raw = clean_text(raw)
    paragraphs = [p for p in raw.split("\n") if p.strip()]
    sections = chunk_by_headers(paragraphs, is_cosmo_pdf_header, source_id)
    return {
        "source": source_label,
        "type": "pdf",
        "totalSections": len(sections),
        "totalChars": sum(len(s.get("body", "")) for s in sections),
        "sections": sections,
    }


def convert_docx(path: Path, source_id: str, source_label: str, header_pred) -> dict:
    from docx import Document
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs]
    sections = chunk_by_headers(paragraphs, header_pred, source_id)
    return {
        "source": source_label,
        "type": "docx",
        "totalSections": len(sections),
        "totalChars": sum(len(s.get("body", "")) for s in sections),
        "sections": sections,
    }


def convert_united_docx(path: Path, source_id: str, source_label: str) -> dict:
    """
    United docx має структуру [короткий-заголовок][довгий-параграф-тексту]....
    Великі body додатково розбиваємо по emoji-під-заголовкам всередині.
    """
    from docx import Document
    doc = Document(str(path))
    paras = [p.text for p in doc.paragraphs if p.text.strip()]

    sections = []
    counter = 0
    i = 0
    while i < len(paras):
        head = paras[i]
        body = ""
        # Якщо наступний параграф довгий — це його тіло
        if i + 1 < len(paras) and len(paras[i + 1]) > 200:
            body = paras[i + 1]
            i += 2
        else:
            i += 1

        # Розбиваємо великі body на під-секції за emoji-заголовками
        sub_sections = split_by_emoji_headers(body, head)
        for sub in sub_sections:
            counter += 1
            slug = re.sub(r"[^a-zа-яёіїє0-9]+", "-", sub["title"].lower(), flags=re.IGNORECASE)[:40].strip("-")
            sections.append({
                "id": f"{source_id}-{counter:02d}-{slug}",
                "title": sub["title"][:120],
                "body": clean_text(sub["body"]),
                "chapter": head[:80],
            })

    return {
        "source": source_label,
        "type": "docx",
        "totalSections": len(sections),
        "totalChars": sum(len(s.get("body", "")) for s in sections),
        "sections": sections,
    }


EMOJI_HEADER_RE = re.compile(r"^([\U0001F300-\U0001FAFF✦◯◇♥◐⚡☉☼☾⊹∞◌🜃🜂🜁🜄🔄🧱🜃🧠🌊🛡⚙️🔹🔮✨🎯🌑])\s+(.+?)(?:\n|$)")


def split_by_emoji_headers(body: str, fallback_title: str) -> list:
    """Розбити body на під-секції за рядками що починаються з emoji."""
    if not body:
        return [{"title": fallback_title, "body": ""}]
    lines = body.split("\n")
    sections = []
    current = {"title": fallback_title, "body": []}
    for line in lines:
        line_s = line.strip()
        # Заголовок — короткий рядок з emoji на початку
        if line_s and len(line_s) < 100 and EMOJI_HEADER_RE.match(line_s):
            if current["body"]:
                current["body"] = "\n".join(current["body"]).strip()
                sections.append(current)
            current = {"title": line_s, "body": []}
        else:
            current["body"].append(line)
    if current["body"]:
        current["body"] = "\n".join(current["body"]).strip()
        sections.append(current)
    return sections if sections else [{"title": fallback_title, "body": body}]


def write_json(name: str, data: dict):
    path = LIBRARY_DIR / name
    with open(path, "w", encoding="utf-8") as f:
        # Мінімізовано — без indent. У runtime indent не потрібен.
        json.dump(data, f, ensure_ascii=False, separators=(",", ":"))
    size_kb = path.stat().st_size / 1024
    item_count = len(data.get("sections", data.get("entries", [])))
    print(f"  wrote {name}: {item_count} items, {size_kb:.1f} KB")


def make_index(books):
    """Компактний індекс — лише id, book, title (без preview-tексту)."""
    index = []
    for book_id, book in books.items():
        for section in book["sections"]:
            # Пропускаємо секції без тіла або з тілом < 50 символів
            body_len = len(section.get("body", ""))
            if body_len < 50:
                continue
            index.append({
                "id": section["id"],
                "book": book_id,
                "title": section["title"][:100],
                "chars": body_len,
            })
    return index


def filter_substantial(book):
    """Прибрати секції з тілом < 50 символів."""
    book["sections"] = [s for s in book["sections"]
                        if len(s.get("body", "")) >= 50]
    book["totalSections"] = len(book["sections"])
    book["totalChars"] = sum(len(s.get("body", "")) for s in book["sections"])
    return book


def main():
    print("Конвертую джерела...")
    books = {}

    if PDF_COSMO.exists():
        print(f"\n→ {PDF_COSMO.name}")
        cosmo = filter_substantial(convert_pdf(PDF_COSMO, "cosmo", "Космоенергетика — повна"))
        write_json("cosmo.json", cosmo)
        books["cosmo"] = cosmo
    else:
        print(f"  ✗ не знайдено: {PDF_COSMO}")

    if DOCX_METH.exists():
        print(f"\n→ {DOCX_METH.name}")
        meth = filter_substantial(convert_docx(DOCX_METH, "meth", "Методична книга Академії", is_methodichka_header))
        write_json("methodichka.json", meth)
        books["meth"] = meth
    else:
        print(f"  ✗ не знайдено: {DOCX_METH}")

    if DOCX_UNITED.exists():
        print(f"\n→ {DOCX_UNITED.name}")
        united = filter_substantial(convert_united_docx(DOCX_UNITED, "united", "Об'єднана база (Космо · Контроль · Потоки · Код)"))
        write_json("united.json", united)
        books["united"] = united
    else:
        print(f"  ✗ не знайдено: {DOCX_UNITED}")

    # Загальний індекс
    print("\n→ index.json (пошуковий)")
    index = make_index(books)
    index_data = {
        "totalEntries": len(index),
        "books": {k: {"label": v["source"], "sections": len(v["sections"])} for k, v in books.items()},
        "entries": index,
    }
    write_json("index.json", index_data)

    print(f"\n✓ Готово. Усього секцій: {len(index)}")


if __name__ == "__main__":
    main()
